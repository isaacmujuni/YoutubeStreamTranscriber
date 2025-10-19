import streamlit as st
import yt_dlp
import speech_recognition as sr
from pydub import AudioSegment
from docx import Document
import tempfile
import os
import threading
import time
from datetime import datetime
import openai

class StreamTranscriber:
    def __init__(self):
        self.recognizer = sr.Recognizer()
        self.is_recording = False
        self.transcription_text = ""
        self.openai_client = None
        
    def setup_openai_client(self, api_key):
        """Setup OpenAI client with API key"""
        try:
            if not api_key:
                st.error("OpenAI API key is required for Whisper API")
                return False
            
            # Test the API key by creating a client
            self.openai_client = openai.OpenAI(api_key=api_key)
            return True
        except Exception as e:
            st.error(f"Error setting up OpenAI client: {str(e)}")
            return False
        
    def extract_audio_from_youtube(self, url):
        """Extract audio from YouTube video/livestream"""
        try:
            # Create temporary directory for download
            temp_dir = tempfile.mkdtemp()
            st.info(f"Created temporary directory: {temp_dir}")
            
            ydl_opts = {
                'format': 'worstaudio[ext=m4a]/worstaudio/bestaudio[abr<=128]/bestaudio[abr<=256]/bestaudio',
                'outtmpl': os.path.join(temp_dir, '%(title)s.%(ext)s'),
                'noplaylist': True,
                'quiet': False,  # Enable verbose output for debugging
                'no_warnings': False,  # Show warnings for debugging
                'extract_flat': False,
                'writeinfojson': False,
                'writesubtitles': False,
                'writeautomaticsub': False,
                'ignoreerrors': True,
                'no_check_certificate': True,
                'prefer_insecure': True,
                'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            }
            
            st.info("Initializing YouTube downloader...")
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                st.info("Extracting video information...")
                # First get info to extract title
                info = ydl.extract_info(url, download=False)
                
                # Handle case where info is None (common with livestreams)
                if info is None:
                    st.warning("Could not extract video info, trying direct download...")
                    title = "Livestream_Recording"
                else:
                    title = info.get('title', 'Unknown')
                    st.info(f"Video title: {title}")
                
                # Clean title for filename
                safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
                
                st.info("Starting audio download...")
                # Download the audio
                ydl.download([url])
                st.info("Download completed, searching for audio files...")
                
                # Find the downloaded file
                downloaded_files = [f for f in os.listdir(temp_dir) if f.endswith(('.wav', '.webm', '.m4a', '.mp3', '.ogg', '.opus'))]
                st.info(f"Found files in temp directory: {os.listdir(temp_dir)}")
                
                if downloaded_files:
                    audio_file = os.path.join(temp_dir, downloaded_files[0])
                    st.success(f"Audio file found: {downloaded_files[0]}")
                    return title, audio_file
                else:
                    st.error("No audio file was downloaded")
                    st.error(f"Files in temp directory: {os.listdir(temp_dir)}")
                    return None, None
                
        except Exception as e:
            st.error(f"Error extracting audio: {str(e)}")
            st.error(f"Error type: {type(e).__name__}")
            import traceback
            st.error(f"Full traceback: {traceback.format_exc()}")
            return None, None
    
    def transcribe_audio_file(self, audio_file_path, use_whisper_api=True, api_key=None):
        """Transcribe audio file to text using OpenAI Whisper API or Google Speech Recognition"""
        try:
            if use_whisper_api:
                return self.transcribe_with_whisper_api(audio_file_path, api_key)
            else:
                return self.transcribe_with_google(audio_file_path)
        except Exception as e:
            st.error(f"Error transcribing audio: {str(e)}")
            return ""
    
    def transcribe_with_whisper_api(self, audio_file_path, api_key):
        """Transcribe using OpenAI Whisper API"""
        try:
            # Load and process audio
            audio = AudioSegment.from_file(audio_file_path)
            
            # Convert to mono and reduce sample rate for better compression
            audio = audio.set_frame_rate(16000).set_channels(1)
            
            # Calculate duration and determine if we need to chunk
            duration_minutes = len(audio) / (1000 * 60)
            max_duration_minutes = 25  # Conservative limit for 25MB
            
            if duration_minutes > max_duration_minutes:
                return self.transcribe_large_audio_chunked(audio, api_key)
            else:
                return self.transcribe_single_chunk(audio, api_key)
                
        except Exception as e:
            st.error(f"Error with Whisper API transcription: {str(e)}")
            return ""
    
    def transcribe_single_chunk(self, audio, api_key):
        """Transcribe a single audio chunk"""
        try:
            temp_mp3 = tempfile.mktemp(suffix='.mp3')
            
            # Export with aggressive compression
            audio.export(temp_mp3, format="mp3", bitrate="32k")
            
            # Check file size
            file_size = os.path.getsize(temp_mp3)
            if file_size > 25 * 1024 * 1024:  # 25MB limit
                st.warning("File still too large, trying more aggressive compression...")
                audio = audio.set_frame_rate(8000)  # Further reduce sample rate
                audio.export(temp_mp3, format="mp3", bitrate="16k")
            
            # Transcribe with Whisper API
            with st.spinner("Transcribing with OpenAI Whisper API..."):
                start_time = time.time()
                
                # Use the client-based approach
                client = openai.OpenAI(api_key=api_key)
                with open(temp_mp3, "rb") as audio_file:
                    transcript = client.audio.transcriptions.create(
                        model="whisper-1",
                        file=audio_file
                    )
                
                elapsed_time = time.time() - start_time
                st.info(f"Transcription completed in {elapsed_time:.1f} seconds")
                
                # Clean up temp file
                if os.path.exists(temp_mp3):
                    os.remove(temp_mp3)
                
                return transcript.text
                
        except Exception as e:
            st.error(f"Error transcribing single chunk: {str(e)}")
            return ""
    
    def transcribe_large_audio_chunked(self, audio, api_key):
        """Transcribe large audio by splitting into chunks"""
        try:
            chunk_duration_ms = 20 * 60 * 1000  # 20 minutes per chunk
            total_chunks = (len(audio) // chunk_duration_ms) + 1
            
            st.info(f"Large audio detected. Splitting into {total_chunks} chunks for processing...")
            
            full_transcript = ""
            
            for i in range(total_chunks):
                start_time = i * chunk_duration_ms
                end_time = min((i + 1) * chunk_duration_ms, len(audio))
                
                if start_time >= len(audio):
                    break
                
                chunk = audio[start_time:end_time]
                
                with st.spinner(f"Processing chunk {i+1}/{total_chunks}..."):
                    chunk_transcript = self.transcribe_single_chunk(chunk, api_key)
                    if chunk_transcript:
                        full_transcript += chunk_transcript + " "
            
            return full_transcript.strip()
            
        except Exception as e:
            st.error(f"Error transcribing large audio: {str(e)}")
            return ""
    
    def transcribe_with_google(self, audio_file_path):
        """Transcribe using Google Speech Recognition (fallback)"""
        try:
            # Load audio file
            audio = AudioSegment.from_file(audio_file_path)
            
            # Convert to WAV format with proper settings for speech recognition
            wav_path = audio_file_path.rsplit('.', 1)[0] + '.wav'
            audio = audio.set_frame_rate(16000).set_channels(1)  # Mono, 16kHz for better recognition
            audio.export(wav_path, format="wav")
            
            # Use speech recognition
            with sr.AudioFile(wav_path) as source:
                # Adjust for ambient noise
                self.recognizer.adjust_for_ambient_noise(source)
                audio_data = self.recognizer.record(source)
                
            # Transcribe using Google Speech Recognition
            text = self.recognizer.recognize_google(audio_data)
            return text
            
        except sr.UnknownValueError:
            st.error("Could not understand the audio")
            return ""
        except sr.RequestError as e:
            st.error(f"Error with speech recognition service: {e}")
            return ""
        except Exception as e:
            st.error(f"Error transcribing audio: {str(e)}")
            return ""
    
    def save_to_word_document(self, text, title, output_path):
        """Save transcribed text to Word document"""
        try:
            doc = Document()
            doc.add_heading(f'Transcription: {title}', 0)
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            # Split text into paragraphs for better formatting
            paragraphs = text.split('. ')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip() + '.')
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            st.error(f"Error saving Word document: {str(e)}")
            return False
    
    def process_youtube_url(self, url, use_whisper_api=True, api_key=None):
        """Main method to process YouTube URL and create transcription"""
        try:
            with st.spinner("Extracting audio from YouTube..."):
                title, audio_file = self.extract_audio_from_youtube(url)
            
            if not title or not audio_file:
                st.error("Failed to extract audio from YouTube. Please check the URL and try again.")
                return False
                
            with st.spinner("Transcribing audio..."):
                transcription = self.transcribe_audio_file(audio_file, use_whisper_api, api_key)
                
            if not transcription:
                st.error("Failed to transcribe audio. Please check your API key and try again.")
                return False
                
            # Save to Word document
            output_filename = f"{title.replace(' ', '_')}_transcription.docx"
            output_path = os.path.join(os.getcwd(), output_filename)
            
            with st.spinner("Saving to Word document..."):
                success = self.save_to_word_document(transcription, title, output_path)
                
            if success:
                st.success(f"Transcription saved as: {output_filename}")
                return True
            else:
                st.error("Failed to save Word document.")
                return False
                
        except Exception as e:
            st.error(f"Unexpected error: {str(e)}")
            return False

def main():
    st.set_page_config(
        page_title="Stream Transcriber",
        page_icon="🎵",
        layout="wide"
    )
    
    st.title("🎵 Stream Transcriber")
    st.markdown("Transcribe YouTube videos and livestreams to Word documents")
    
    # Initialize transcriber in session state to avoid reinitialization
    if 'transcriber' not in st.session_state:
        st.session_state.transcriber = StreamTranscriber()
    
    transcriber = st.session_state.transcriber
    
    # Input section
    st.header("📝 Enter YouTube URL")
    
    # URL input with session state
    if 'url' not in st.session_state:
        st.session_state.url = ""
    
    url = st.text_input("YouTube Video or Livestream URL:", 
                       value=st.session_state.url,
                       placeholder="https://www.youtube.com/watch?v=...")
    
    # Update session state when URL changes
    if url != st.session_state.url:
        st.session_state.url = url
    
    # Advanced options
    st.header("⚙️ Transcription Options")
    
    # API Key input with session state
    if 'api_key' not in st.session_state:
        st.session_state.api_key = ""
    
    api_key = st.text_input("OpenAI API Key:", type="password", 
                           value=st.session_state.api_key,
                           help="Get your API key from https://platform.openai.com/api-keys")
    
    # Update session state when API key changes
    if api_key != st.session_state.api_key:
        st.session_state.api_key = api_key
    
    col1, col2 = st.columns(2)
    
    with col1:
        use_whisper_api = st.checkbox("Use OpenAI Whisper API (Recommended)", value=True, 
                                     help="Faster, more accurate, and supports multiple languages")
    
    with col2:
        if not use_whisper_api:
            st.info("Using Google Speech Recognition (requires internet)")
        else:
            st.info("Using OpenAI Whisper API (requires API key)")
    
    col1, col2 = st.columns([1, 4])
    
    with col1:
        transcribe_button = st.button("🎯 Transcribe", type="primary")
    
    with col2:
        if transcribe_button and url:
            if use_whisper_api and not api_key:
                st.error("Please enter your OpenAI API key to use Whisper API")
            elif transcriber.process_youtube_url(url, use_whisper_api, api_key):
                st.balloons()
            else:
                st.error("Transcription failed. Please check the URL and API key, then try again.")
    
    # Instructions
    st.header("📋 How to Use")
    st.markdown("""
    1. **Get OpenAI API Key** - Sign up at [OpenAI Platform](https://platform.openai.com/api-keys)
    2. **Copy a YouTube URL** - Paste any YouTube video or livestream URL
    3. **Enter API Key** - Paste your OpenAI API key (required for Whisper API)
    4. **Click Transcribe** - The app will extract audio and transcribe it
    5. **Download** - Get your Word document with the transcription
    
    **Note:** OpenAI Whisper API is faster and more accurate. Google Speech Recognition is free but requires internet.
    """)
    
    # Features
    st.header("✨ Features")
    st.markdown("""
    - 🎥 **YouTube Support** - Works with videos and livestreams
    - 🎵 **Audio Extraction** - Automatically extracts audio from video
    - 🤖 **OpenAI Whisper API** - Fast, accurate speech recognition with multi-language support
    - 🌐 **Google Speech Recognition** - Free fallback option requiring internet connection
    - 📄 **Word Documents** - Saves transcriptions in .docx format
    - ⚡ **Fast Processing** - Quick transcription and document generation
    - 🔒 **Secure** - API keys are not stored, processed locally only
    - 💰 **Cost Effective** - Pay only for what you use (~$0.006 per minute)
    """)

if __name__ == "__main__":
    main()
